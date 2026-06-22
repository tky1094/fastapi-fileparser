import io

from docx import Document

from app.parsers.base import BaseParser, ParseResult, ProgressCallback


class DocxParser(BaseParser):
    def supported_mimetypes(self) -> set[str]:
        return {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

    async def parse(
        self,
        content: bytes,
        mime_type: str,
        filename: str,
        progress_callback: ProgressCallback | None = None,
    ) -> ParseResult:
        await self._notify(progress_callback, "parsing", "Word文書を解析中...")

        doc = Document(io.BytesIO(content))

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract table text
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)

        all_parts: list[str] = []
        if paragraphs:
            all_parts.append("\n".join(paragraphs))
        if table_texts:
            all_parts.append("\n".join(table_texts))

        text = "\n\n".join(all_parts)

        return ParseResult(
            text=text,
            content_type=mime_type,
            parser_used="docx",
            metadata={
                "paragraph_count": str(len(paragraphs)),
                "table_count": str(len(doc.tables)),
            },
        )
