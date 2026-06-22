import io

import pytest
from docx import Document

from app.parsers.docx import DocxParser


def create_test_docx(paragraphs: list[str], table_data: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def parser() -> DocxParser:
    return DocxParser()


async def test_parse_paragraphs(parser: DocxParser) -> None:
    content = create_test_docx(["Hello", "World", "Test"])
    result = await parser.parse(
        content,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "test.docx",
    )
    assert "Hello" in result.text
    assert "World" in result.text
    assert result.parser_used == "docx"
    assert result.metadata["paragraph_count"] == "3"


async def test_parse_with_table(parser: DocxParser) -> None:
    content = create_test_docx(
        ["Intro"],
        table_data=[["A", "B"], ["C", "D"]],
    )
    result = await parser.parse(
        content,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "test.docx",
    )
    assert "Intro" in result.text
    assert "A" in result.text
    assert result.metadata["table_count"] == "1"


async def test_parse_empty_docx(parser: DocxParser) -> None:
    content = create_test_docx([])
    result = await parser.parse(
        content,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "empty.docx",
    )
    assert result.text == ""


async def test_supported_mimetypes(parser: DocxParser) -> None:
    mimes = parser.supported_mimetypes()
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in mimes
